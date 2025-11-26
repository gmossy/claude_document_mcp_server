#!/usr/bin/env python3
"""
Generate random test files for testing document upload/download endpoints.

Creates:
- Word documents (.docx)
- PDF files (.pdf)
- Text files (.txt)
- Excel files (.xlsx)
- CUE language files (.cue)

All files are random sizes up to 5 pages.
"""

import random
import string
from pathlib import Path
from datetime import datetime

# Word document generation
try:
    from docx import Document
    from docx.shared import Pt, Inches
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False
    print("⚠️  python-docx not available, skipping Word files")

# PDF generation
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.pdfgen import canvas
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️  reportlab not available, skipping PDF files")

# Excel generation
try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("⚠️  openpyxl not available, skipping Excel files")

def generate_random_text(words=100):
    """Generate random text with specified word count."""
    word_list = [
        "document", "test", "sample", "content", "data", "information",
        "system", "application", "server", "client", "database", "network",
        "security", "performance", "optimization", "development", "testing",
        "deployment", "configuration", "management", "analysis", "design",
        "implementation", "architecture", "framework", "library", "module",
        "function", "method", "class", "object", "variable", "constant",
        "algorithm", "structure", "pattern", "protocol", "interface", "api"
    ]
    sentences = []
    for _ in range(words // 10):  # ~10 words per sentence
        sentence = ' '.join(random.choices(word_list, k=random.randint(8, 15)))
        sentence = sentence.capitalize() + '.'
        sentences.append(sentence)
    return ' '.join(sentences)

def generate_word_document(output_path, pages=3):
    """Generate a Word document with random content."""
    if not WORD_AVAILABLE:
        return False
    
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Test Document - {datetime.now().strftime("%Y%m%d_%H%M%S")}', 0)
    
    # Generate content for specified pages
    # Approx 500 words per page
    words_per_page = 500
    total_words = pages * words_per_page
    
    # Add paragraphs
    text = generate_random_text(total_words)
    paragraphs = text.split('. ')
    
    for para_text in paragraphs:
        if para_text.strip():
            p = doc.add_paragraph(para_text.strip() + '.')
            p.style.font.size = Pt(11)
    
    # Add some formatting
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(generate_random_text(100))
    
    doc.add_heading('Details', level=1)
    doc.add_paragraph(generate_random_text(200))
    
    doc.save(output_path)
    return True

def generate_pdf_file(output_path, pages=3):
    """Generate a PDF file with random content."""
    if not PDF_AVAILABLE:
        return False
    
    c = canvas.Canvas(str(output_path), pagesize=letter)
    width, height = letter
    
    # Generate content
    text = generate_random_text(pages * 500)
    lines = text.split('. ')
    
    y_position = height - 50
    line_height = 14
    lines_per_page = int((height - 100) / line_height)
    
    page_num = 1
    line_count = 0
    
    for line in lines:
        if line_count >= lines_per_page:
            c.showPage()
            page_num += 1
            y_position = height - 50
            line_count = 0
        
        if line.strip():
            c.drawString(50, y_position, line.strip()[:80] + '.')
            y_position -= line_height
            line_count += 1
    
    # Add page numbers (already added during page creation)
    c.save()
    return True

def generate_text_file(output_path, pages=3):
    """Generate a plain text file with random content."""
    # Approx 500 words per page
    words_per_page = 500
    total_words = pages * words_per_page
    
    text = generate_random_text(total_words)
    
    # Add some structure
    content = f"""Test Document - {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
{'=' * 60}

{text}

{'=' * 60}
Summary
{'=' * 60}

{generate_random_text(100)}

{'=' * 60}
Details
{'=' * 60}

{generate_random_text(200)}

{'=' * 60}
Conclusion
{'=' * 60}

{generate_random_text(150)}
"""
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    return True

def generate_excel_file(output_path, pages=3):
    """Generate an Excel file with random data."""
    if not EXCEL_AVAILABLE:
        return False
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Test Data"
    
    # Headers
    headers = ["ID", "Name", "Category", "Value", "Date", "Status", "Notes"]
    ws.append(headers)
    
    # Style headers
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center')
    
    # Generate data rows
    # Approx 50 rows per page
    rows_per_page = 50
    total_rows = pages * rows_per_page
    
    categories = ["Type A", "Type B", "Type C", "Type D", "Type E"]
    statuses = ["Active", "Pending", "Completed", "Archived"]
    
    for i in range(1, total_rows + 1):
        row = [
            f"ID-{i:04d}",
            f"Item {i}",
            random.choice(categories),
            round(random.uniform(10.0, 1000.0), 2),
            datetime.now().strftime("%Y-%m-%d"),
            random.choice(statuses),
            generate_random_text(20)
        ]
        ws.append(row)
    
    # Add summary sheet
    ws2 = wb.create_sheet("Summary")
    ws2.append(["Metric", "Value"])
    ws2.append(["Total Items", total_rows])
    ws2.append(["Average Value", round(random.uniform(100.0, 500.0), 2)])
    ws2.append(["Generated", datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
    
    wb.save(output_path)
    return True

def generate_cue_file(output_path, pages=3):
    """Generate a CUE language file (configuration file format)."""
    # CUE files are typically smaller, so we'll generate multiple sections
    sections = pages * 2
    
    words = generate_random_text(5).split()
    name_word = words[0] if words else "test"
    desc_words = generate_random_text(20).split()[:10]
    description = ' '.join(desc_words) if desc_words else "Test configuration"
    
    content = f"""// CUE Configuration File
// Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
// Test document for document management system

package test

import (
    "encoding/json"
    "strings"
)

// Configuration structure
config: {{
    name: "{name_word}"
    version: "1.0.0"
    description: "{description}"
    
    // Server configuration
    server: {{
        host: string | *"localhost"
        port: int | *8080
        timeout: int | *30
    }}
    
    // Database configuration
    database: {{
        type: "sqlite" | "postgres" | "mysql"
        connection: string
        poolSize: int | *10
    }}
    
    // Feature flags
    features: {{
        caching: bool | *true
        logging: bool | *true
        metrics: bool | *false
    }}
}}

"""
    
    # Add multiple sections
    for i in range(sections):
        section_name = f"section_{i+1}"
        name_words = generate_random_text(3).split()
        section_name_val = name_words[0] if name_words else f"section{i+1}"
        value3_words = generate_random_text(10).split()[:5]
        value3_text = ' '.join(value3_words) if value3_words else "test value"
        
        content += f"""
// {section_name} configuration
{section_name}: {{
    id: "{random.randint(1000, 9999)}"
    name: "{section_name_val}"
    enabled: bool | *true
    settings: {{
        value1: {random.randint(1, 100)}
        value2: {random.uniform(1.0, 100.0)}
        value3: "{value3_text}"
    }}
}}
"""
    
    content += """
// Validation rules
validation: {
    minLength: int | *5
    maxLength: int | *100
    pattern: string | *"^[a-zA-Z0-9]+$"
}

// Export configuration
export: {
    format: "json" | "yaml" | "toml"
    pretty: bool | *true
}
"""
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    return True

def main():
    """Generate all test files."""
    # Determine correct path (host or container)
    import os
    host_path = Path("/Users/glennmossy/dpg-ai-projects/claude_document_mcp_server/testfiles")
    container_path = Path("/data/testfiles")
    
    # Check if we're in Docker (check for /data mount)
    if Path("/data").exists() and Path("/data/documents.db").exists():
        # We're in Docker, use /data/testfiles (mounted to host ./data/testfiles)
        testfiles_dir = container_path
        print(f"Running in Docker container, using: {testfiles_dir}")
    else:
        # We're on host, use host path
        testfiles_dir = host_path
        print(f"Running on host, using: {testfiles_dir}")
    
    testfiles_dir.mkdir(parents=True, exist_ok=True)
    
    print("=" * 60)
    print("Generating Test Files")
    print("=" * 60)
    print(f"Output directory: {testfiles_dir}\n")
    
    files_created = []
    
    # Generate Word documents (2-5 pages each)
    if WORD_AVAILABLE:
        for i in range(3):
            pages = random.randint(2, 5)
            filename = f"Word_Test_{i+1}_{pages}pages_{random.randint(1000, 9999)}.docx"
            filepath = testfiles_dir / filename
            if generate_word_document(filepath, pages):
                size = filepath.stat().st_size
                files_created.append(("Word", filename, pages, size))
                print(f"✅ Created: {filename} ({pages} pages, {size:,} bytes)")
    
    # Generate PDF files (2-5 pages each)
    if PDF_AVAILABLE:
        for i in range(3):
            pages = random.randint(2, 5)
            filename = f"PDF_Test_{i+1}_{pages}pages_{random.randint(1000, 9999)}.pdf"
            filepath = testfiles_dir / filename
            if generate_pdf_file(filepath, pages):
                size = filepath.stat().st_size
                files_created.append(("PDF", filename, pages, size))
                print(f"✅ Created: {filename} ({pages} pages, {size:,} bytes)")
    
    # Generate text files (2-5 pages each)
    for i in range(3):
        pages = random.randint(2, 5)
        filename = f"Text_Test_{i+1}_{pages}pages_{random.randint(1000, 9999)}.txt"
        filepath = testfiles_dir / filename
        if generate_text_file(filepath, pages):
            size = filepath.stat().st_size
            files_created.append(("Text", filename, pages, size))
            print(f"✅ Created: {filename} ({pages} pages, {size:,} bytes)")
    
    # Generate Excel files (2-5 pages each)
    if EXCEL_AVAILABLE:
        for i in range(3):
            pages = random.randint(2, 5)
            filename = f"Excel_Test_{i+1}_{pages}pages_{random.randint(1000, 9999)}.xlsx"
            filepath = testfiles_dir / filename
            if generate_excel_file(filepath, pages):
                size = filepath.stat().st_size
                files_created.append(("Excel", filename, pages, size))
                print(f"✅ Created: {filename} ({pages} pages, {size:,} bytes)")
    
    # Generate CUE files
    for i in range(3):
        pages = random.randint(2, 5)
        filename = f"CUE_Test_{i+1}_{pages}pages_{random.randint(1000, 9999)}.cue"
        filepath = testfiles_dir / filename
        if generate_cue_file(filepath, pages):
            size = filepath.stat().st_size
            files_created.append(("CUE", filename, pages, size))
            print(f"✅ Created: {filename} ({pages} pages, {size:,} bytes)")
    
    print("\n" + "=" * 60)
    print("Summary")
    print("=" * 60)
    print(f"\nTotal files created: {len(files_created)}")
    print(f"\nBreakdown by type:")
    for file_type in ["Word", "PDF", "Text", "Excel", "CUE"]:
        count = sum(1 for f in files_created if f[0] == file_type)
        if count > 0:
            print(f"  {file_type}: {count} files")
    
    total_size = sum(f[3] for f in files_created)
    print(f"\nTotal size: {total_size:,} bytes ({total_size / 1024:.2f} KB)")
    print(f"\nFiles saved to: {testfiles_dir}")
    print("=" * 60)

if __name__ == "__main__":
    main()

