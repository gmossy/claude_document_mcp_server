import json
from pathlib import Path

import pytest

import document_mcp_server as server
from document_parsers import extract_text_from_file


@pytest.fixture(autouse=True)
def clean_test_db(tmp_path, monkeypatch):
    """
    Use a temporary SQLite database for tests so we don't touch the real one.
    """
    test_db = tmp_path / "test_documents.db"
    monkeypatch.setattr(server, "DATABASE_PATH", test_db)

    # (Re)initialise the database for each test
    server.init_database()
    yield


@pytest.mark.asyncio
async def test_document_create_and_get_roundtrip():
    """
    Basic sanity test: create a document and fetch it back via the MCP tools.
    """
    # Create a document
    create_input = server.CreateDocumentInput(
        title="Test Doc",
        content="Hello world from test.",
        tags=["test"],
    )

    create_raw = await server.document_create(create_input)
    create_resp = json.loads(create_raw)

    assert create_resp["success"] is True
    doc_id = create_resp["document_id"]
    assert isinstance(doc_id, str) and doc_id.startswith("doc_")

    # Get the document back in JSON format
    get_input = server.GetDocumentInput(
        document_id=doc_id,
        include_content=True,
        include_versions=True,
        response_format=server.ResponseFormat.JSON,
    )

    get_raw = await server.document_get(get_input)
    get_resp = json.loads(get_raw)

    assert get_resp["id"] == doc_id
    assert get_resp["title"] == "Test Doc"
    assert get_resp["content"] == "Hello world from test."
    assert get_resp["tags"] == ["test"]


@pytest.mark.asyncio
async def test_large_word_document_ingest(tmp_path, monkeypatch):
    """
    Create a ~10-page Word document, extract its text, and store it via the MCP tools.
    This exercises the pipeline on a realistically large document.
    """
    # Generate a long Word document (~10 pages worth of text)
    from docx import Document

    docx_path = tmp_path / "long_test.docx"
    document = Document()
    document.add_heading("Long Test Document", 0)

    # Rough heuristic: many paragraphs of text to span ~10 pages
    paragraph_text = (
        "This is a long test paragraph intended to help validate large Word document "
        "processing in the Claude document MCP server. It contains multiple sentences "
        "and is repeated many times to generate enough content for a multi-page file. "
        "The content itself is not important; only the size and structure matter.\n"
    )
    for _ in range(200):  # 200 paragraphs is plenty for ~10 pages of text
        document.add_paragraph(paragraph_text)

    document.save(docx_path)

    # Extract text using the existing parser util
    text, metadata = extract_text_from_file(docx_path)
    assert len(text) > 10_000  # sanity check: it's really "large"
    assert metadata.get("format") == "docx"

    # Point the server at a fresh temporary DB
    test_db = tmp_path / "large_doc_test.db"
    monkeypatch.setattr(server, "DATABASE_PATH", test_db)
    server.init_database()

    # Create a document in the MCP server using the extracted text
    create_input = server.CreateDocumentInput(
        title="Long Word Doc",
        content=text,
        tags=["word", "large", "test"],
    )
    create_raw = await server.document_create(create_input)
    create_resp = json.loads(create_raw)

    assert create_resp["success"] is True
    long_doc_id = create_resp["document_id"]

    # Fetch it back and verify core fields
    get_input = server.GetDocumentInput(
        document_id=long_doc_id,
        include_content=True,
        include_versions=False,
        response_format=server.ResponseFormat.JSON,
    )
    get_raw = await server.document_get(get_input)
    get_resp = json.loads(get_raw)

    assert get_resp["id"] == long_doc_id
    assert get_resp["title"] == "Long Word Doc"
    assert isinstance(get_resp["content"], str)
    # content should be substantial
    assert len(get_resp["content"]) >= len(text) * 0.9


